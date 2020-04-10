from data import db_session
from data.users import User
from data.register_form import RegisterForm
from data.login_form import LoginForm
from data.tournament_form import TournamentForm
from data.tournaments import Tournament
from data.tasks import Task
from data.results import UserResult
from data.attempts import Attempt, AllAttempt
from data.get_task_form import GetTaskForm
from flask import Flask, url_for, request, render_template, redirect, abort, jsonify, make_response
from data.send_task_form import SendTaskForm
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
import datetime
from docx import Document
import random
import os


app = Flask(__name__)
login_manager = LoginManager()
login_manager.init_app(app)
app.config['SECRET_KEY'] = 'tournament_of_games_secret_key'


@login_manager.user_loader
def load_user(user_id):
    session = db_session.create_session()
    return session.query(User).get(user_id)


@app.route('/')
def index():
    session = db_session.create_session()
    tour = session.query(Tournament).first()
    tour_is_started = session.query(Tournament).first() is not None
    data = dict()
    data["tour_is_started"] = tour_is_started
    if not tour_is_started or not current_user.is_authenticated:
        return render_template("index.html", data=data)
    data["type"] = tour.type
    if current_user.id > 2:
        res = session.query(UserResult).filter(UserResult.user == current_user.id).first()
        data['flag'] = res is None
        if data['flag']:
            return render_template("index.html", data=data)
        now_tasks = res.tasks.split("@")
        if res.tasks == "":
            data['get_task'] = True
            return render_template("index.html", data=data)
        if len(now_tasks) < 2:
            data['get_task'] = True
        data['now_tasks'] = []
        for now_task in now_tasks:
            text = session.query(Task).filter(Task.num == now_task).first().text
            data['now_tasks'].append((now_task, text))
    else:
        attempts = session.query(Attempt).all()
        data['attempts'] = []
        for attempt in attempts:
            user = session.query(User).filter(User.id == attempt.user).first()
            name = user.name + " " + user.surname
            data['attempts'].append((name, attempt.num, attempt.answer, attempt.image, attempt.user))
    return render_template("index.html", data=data)


@app.route('/register', methods=['POST', 'GET'])
def register():
    session = db_session.create_session()
    form = RegisterForm()
    if form.validate_on_submit():
        user = User(
            surname=form.surname.data,
            name=form.name.data,
            num_of_class=form.num_of_class.data,
            email=form.login.data
        )
        user.set_password(form.password.data)
        session.add(user)
        session.commit()
        login_user(user, remember=False)
        return redirect("/")
    return render_template('register.html', title='Регистрация', form=form)


@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        session = db_session.create_session()
        user = session.query(User).filter(User.email == form.login.data).first()
        if user and user.check_password(form.password.data):
            login_user(user, remember=form.remember_me.data)
            return redirect("/")
        return render_template('login.html',
                               message="Неправильный адрес почты или пароль",
                               form=form)
    return render_template('login.html', form=form)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect("/")


@app.route('/create_tournament', methods=['GET', 'POST'])
@login_required
def creating_tournament():
    if current_user.id != 1 and current_user.id != 2:
        return redirect("/")
    form = TournamentForm()
    form.type_of_tournament.choices = [('Домино', 'Домино'), ('Пенальти', 'Пенальти')]
    if form.validate_on_submit():
        session = db_session.create_session()

        dop = session.query(Tournament).first()
        if dop is not None:
            session.delete(dop)
            session.commit()
        dop = session.query(Task).all()
        for task in dop:
            session.delete(task)
            session.commit()

        text_file = form.text.data
        text_filename = text_file.filename
        text_file.save(os.path.join(text_filename))

        answers_file = form.answers.data
        answers_filename = answers_file.filename
        answers_file.save(os.path.join(answers_filename))
        tour = Tournament(
            name=form.name.data,
            type=form.type_of_tournament.data,
            time=form.time.data,
            text=text_filename,
            answers=answers_filename
        )
        session.add(tour)
        session.commit()
        creating_tasks(int(form.num_of_tasks.data))
        return redirect('/')
    return render_template('creating_tournament.html', form=form)


@app.route('/delete_tournament', methods=['GET', 'POST'])
@login_required
def deleting_tournament():
    if current_user.id > 2:
        return redirect("/")
    session = db_session.create_session()
    tour = session.query(Tournament).first()
    session.delete(tour)
    session.commit()

    tasks = session.query(Task).all()
    for task in tasks:
        session.delete(task)
        session.commit()

    results = session.query(UserResult).all()
    for res in results:
        session.delete(res)
        session.commit()

    attempts = session.query(Attempt).all()
    for attempt in attempts:
        session.delete(attempt)
        session.commit()
    return redirect("/")


def creating_tasks(cnt_task):
    session = db_session.create_session()
    tour = session.query(Tournament).first()
    text_file = Document(tour.text)
    for para in text_file.paragraphs:
        dop = para.text.strip()
        if dop != "":
            ind = dop.find('.')
            num = dop[1:ind].strip()
            text = dop[ind + 1:].strip()
            task = Task(
                num=num,
                text=text,
                cnt=cnt_task
            )
            session.add(task)
            session.commit()
    os.remove(tour.text)

    answers_file = Document(tour.answers)
    for para in answers_file.paragraphs:
        dop = para.text.strip()
        if dop != "":
            ind = dop.find('.')
            num = dop[1:ind].strip()
            answer = dop[ind + 1:].strip()
            task = session.query(Task).filter(Task.num == num).first()
            try:
                task.answer = answer
                session.commit()
            except AttributeError:
                pass
    os.remove(tour.answers)


@app.route("/results")
def results():
    session = db_session.create_session()
    tour = session.query(Tournament).first()
    if tour is None:
        return redirect("/")
    tasks = session.query(Task).all()
    headers = ["Место", "Участник"]
    for task in tasks:
        dop_num = " ".join(task.num.split('–'))
        headers.append(dop_num)
    headers.append("Сумма")
    results = session.query(UserResult).all()
    data = []
    for res in results:
        user = session.query(User).filter(User.id == res.user).first()
        dop = [(user.name + " " + user.surname, "white")]
        sum_points = 0
        points = res.cnt_points.split("@")
        attempts = string_to_dict(res.attempts)
        for elem in points:
            num, point = elem.split("$")
            if int(attempts[num]) == 0:
                dop.append(("", "white"))
            elif (int(attempts[num]) == 2 and int(point) == 0 or int(attempts[num]) == 1
                    and num == "0–0" and int(point) == 0):
                dop.append((point, "yellow"))
            else:
                dop.append((point, "white"))
            sum_points += int(point)
        dop.append((str(sum_points), "green"))
        data.append(dop)
    data.sort(key=lambda x: -int(x[-1][0]))
    dop = max(get_now_time(), datetime.datetime.now() - datetime.datetime.now())
    time = (str(dop.seconds // 3600), str(dop.seconds % 3600 // 60), str(dop.seconds % 60))
    return render_template("results.html", headers=headers, data=data, time=time)


def dict_to_string(dop):
    res = []
    for key in dop.keys():
        res.append(key + "$" + dop[key])
    return "@".join(res)


def string_to_dict(dop):
    dop = dop.split("@")
    res = dict()
    for elem in dop:
        key, val = elem.split("$")
        res[key] = val
    return res


@app.route("/get_start_tasks")
@login_required
def getting_start_task():
    if current_user.id < 3:
        return redirect("/")
    session = db_session.create_session()
    res = session.query(UserResult).filter(UserResult.user == current_user.id).first()
    if res is not None:
        return redirect("/")
    tasks = session.query(Task).all()
    tour = session.query(Tournament).first()
    if tour.type == "Пенальти":
        res = UserResult(
            user=current_user.id,
            cnt_points="@".join([str(i) + "$0" for i in range(1, 17)]),
            attempts="@".join([str(i) + "$0" for i in range(1, 17)]),
            cnt_tasks=16,
            tasks="@".join([str(i) for i in range(1, 17)])
        )
    else:
        dop = []
        for task in tasks:
            a1, a2 = map(int, task.num.split('–'))
            if a1 + a2 < 8 and a1 + a2 != 0 and task.cnt > 0:
                dop.append(task.num)
        dop2 = []
        while True:
            task1, task2 = random.choice(dop), random.choice(dop)
            if task1 == task2:
                continue
            a1, a2 = map(int, task1.split('–'))
            a3, a4 = map(int, task2.split('–'))
            if a1 + a2 + a3 + a4 <= 8:
                dop2.append(task1)
                obj_task = session.query(Task).filter(Task.num == task1).first()
                obj_task.cnt -= 1
                session.commit()
                dop2.append(task2)
                obj_task = session.query(Task).filter(Task.num == task2).first()
                obj_task.cnt -= 1
                session.commit()
                break
        dop = []
        for task in tasks:
            dop.append(task.num)
        res = UserResult(
            user=current_user.id,
            cnt_points="@".join([num + "$0" for num in dop]),
            attempts="@".join([num + "$0" for num in dop]),
            cnt_tasks=2,
            tasks="@".join([num for num in dop2])
        )
    session.add(res)
    session.commit()
    return redirect("/")


@app.route('/delete_task/<num_task>', methods=['GET', 'POST'])
@login_required
def delete_task(num_task):
    if datetime.datetime.now() - datetime.datetime.now() > get_now_time():
        return redirect("/")
    session = db_session.create_session()
    task = session.query(Task).filter(Task.num == num_task).first()
    task.cnt += 1
    session.commit()
    res = session.query(UserResult).filter(UserResult.user == current_user.id).first()
    now_tasks = res.tasks.split("@")
    new_now_tasks = []
    for item in now_tasks:
        if item != num_task:
            new_now_tasks.append(item)
    res.tasks = "@".join(new_now_tasks)
    attempts = string_to_dict(res.attempts)
    attempts[num_task] = str(int(attempts[num_task]) + 1)
    res.attempts = dict_to_string(attempts)
    session.commit()
    return redirect("/")


@app.route('/send_task/<num_task>', methods=['GET', 'POST'])
@login_required
def send_task(num_task):
    if datetime.datetime.now() - datetime.datetime.now() > get_now_time():
        return redirect("/")
    form = SendTaskForm()
    form.num.data = num_task
    session = db_session.create_session()
    res = session.query(UserResult).filter(UserResult.user == current_user.id).first()
    if num_task not in res.tasks:
        return redirect("/")
    if form.validate_on_submit():
        session = db_session.create_session()
        task = session.query(Task).filter(Task.num == num_task).first()
        if task.answer is not None:
            res = session.query(UserResult).filter(UserResult.user == current_user.id).first()
            points = string_to_dict(res.cnt_points)
            if task.answer == form.answer.data.strip():
                tour = session.query(Tournament).first()
                if tour.type == "Домино":
                    points = string_to_dict(res.cnt_points)
                    attempts = string_to_dict(res.attempts)
                    a1, a2 = map(int, num_task.split('–'))
                    if a1 + a2 == 0:
                        a1 = 10
                    if int(attempts[num_task]) == 0:
                        points[num_task] = str(a1 + a2)
                    elif int(attempts[num_task]) == 1:
                        points[num_task] == str(max(a1, a2))
                elif tour.type == "Пенальти":
                    now_point = 10
                    all_res = session.query(UserResult).all()
                    for dop_res in all_res:
                        points = string_to_dict(dop_res.cnt_points)
                        if int(points[num_task]) != 0:
                            now_point = min(now_point, int(points[num_task]))
                    point = max(now_point - 1, 5)
                    points = string_to_dict(res.cnt_points)
                    attempts = string_to_dict(res.attempts)
                    if int(attempts[num_task]) == 0:
                        points[num_task] = str(point)
            now_tasks = res.tasks.split("@")
            new_now_tasks = []
            for item in now_tasks:
                if item != num_task:
                    new_now_tasks.append(item)
            res.tasks = "@".join(new_now_tasks)
            attempts = string_to_dict(res.attempts)
            attempts[num_task] = str(int(attempts[num_task]) + 1)
            res.attempts = dict_to_string(attempts)
            res.cnt_points = dict_to_string(points)
            task.cnt += 1
            session.commit()
            dop_attempt = AllAttempt(
                num=num_task,
                user=current_user.id,
                answer=form.answer.data
            )
            session.add(dop_attempt)
            session.commit()
        else:
            res = session.query(UserResult).filter(UserResult.user == current_user.id).first()
            now_tasks = res.tasks.split("@")
            new_now_tasks = []
            for item in now_tasks:
                if item != num_task:
                    new_now_tasks.append(item)
            res.tasks = "@".join(new_now_tasks)
            attempts = string_to_dict(res.attempts)
            attempts[num_task] = str(int(attempts[num_task]) + 1)
            res.attempts = dict_to_string(attempts)
            session.commit()

            if form.image.data is not None:
                image_file = form.image.data
                image_filename = "static/img/" + str(current_user.id) + "_" + str(num_task)
                image_file.save(os.path.join(image_filename))
            else:
                image_filename = ""

            attempt = Attempt(
                num=num_task,
                user=current_user.id,
                answer=form.answer.data,
                image=image_filename
            )
            session.add(attempt)
            task.cnt += 1
            session.commit()
        return redirect("/")
    return render_template('sending_task.html', form=form)


@app.route('/ignore_attempt/<int:user_id>/<num_task>', methods=['GET', 'POST'])
@login_required
def ignore_attempt(user_id, num_task):
    session = db_session.create_session()
    attempt = session.query(Attempt).filter(Attempt.user == user_id, Attempt.num == num_task).first()
    session.delete(attempt)
    session.commit()
    return redirect("/")


@app.route('/success_attempt/<int:user_id>/<num_task>', methods=['GET', 'POST'])
@login_required
def success_attempt(user_id, num_task):
    session = db_session.create_session()
    attempt = session.query(Attempt).filter(Attempt.user == user_id, Attempt.num == num_task).first()
    session.delete(attempt)
    session.commit()

    res = session.query(UserResult).filter(UserResult.user == user_id).first()
    points = string_to_dict(res.cnt_points)
    tour = session.query(Tournament).first()
    if tour.type == "Домино":
        points = string_to_dict(res.cnt_points)
        attempts = string_to_dict(res.attempts)
        a1, a2 = map(int, num_task.split('–'))
        if a1 + a2 == 0:
            a1 = 10
        if int(attempts[num_task]) == 1:
            points[num_task] = str(a1 + a2)
        elif int(attempts[num_task]) == 2:
            points[num_task] == str(max(a1, a2))
    elif tour.type == "Пенальти":
        now_point = 10
        all_res = session.query(UserResult).all()
        for dop_res in all_res:
            points = string_to_dict(dop_res.cnt_points)
            if int(points[num_task]) != 0:
                now_point = min(now_point, int(points[num_task]))
        point = max(now_point - 1, 5)
        points = string_to_dict(res.cnt_points)
        attempts = string_to_dict(res.attempts)
        if int(attempts[num_task]) == 0:
            points[num_task] = str(point)
    res.cnt_points = dict_to_string(points)
    session.commit()
    return redirect("/")


@app.route("/get_task", methods=['GET', 'POST'])
@login_required
def getting_task():
    if datetime.datetime.now() - datetime.datetime.now() > get_now_time():
        return redirect("/")
    if current_user.id < 3:
        return redirect("/")
    session = db_session.create_session()
    form = GetTaskForm()
    form.num_task.choices = []
    tasks = session.query(Task).all()
    res = session.query(UserResult).filter(UserResult.user == current_user.id).first()
    attempts = string_to_dict(res.attempts)
    points = string_to_dict(res.cnt_points)
    if res.tasks.count("@") == 1:
        return redirect("/")
    for task in tasks:
        if ((int(attempts[task.num]) == 0 and task.num == "0–0" or int(attempts[task.num]) < 2 and
             task.num != "0–0") and
                task.num not in res.tasks and task.cnt > 0 and int(points[task.num]) == 0):
            form.num_task.choices.append((task.num, task.num))
    if form.validate_on_submit():
        num_task = form.num_task.data
        task = session.query(Task).filter(Task.num == num_task).first()
        if task.cnt > 0:
            if res.tasks.strip() == "":
                res.tasks = num_task
            else:
                res.tasks = res.tasks + "@" + num_task
            task.cnt -= 1
        else:
            return render_template("success.html",
                                   message="Данная задача уже не доступна. Попробуйте позже")
        session.commit()
        return redirect('/')
    return render_template('getting_task.html', form=form)


def get_now_time():
    session = db_session.create_session()
    tour = session.query(Tournament).first()
    dop = datetime.timedelta(seconds=tour.time.second + tour.time.minute * 60 + tour.time.hour * 3600)
    time = (tour.start_date + dop) - datetime.datetime.now()
    return time


def main():
    db_session.global_init("db/data_base.db")
    app.run(host="https://tournamentofgames.herokuapp.com")

if __name__ == "__main__":
    main()